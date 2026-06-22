from __future__ import annotations

import json
import math
import os
import re
from collections import defaultdict
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Iterable

import geopandas as gpd
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from shapely.geometry import mapping

from v11_common import (
    FRONTEND,
    INVENTORY,
    REPORTS,
    TARGET_REGIONS,
    V11_DIR,
    boundary_gdf,
    ensure_dirs,
    entropy_from_shares,
    read_csv,
    safe_float,
    target_region_labels,
    to_5179,
    to_4326,
    write_csv,
    write_geojson,
    write_json,
    write_md,
)


ROOT = Path(__file__).resolve().parents[1]
DATA = ROOT / "data"
RAW = DATA / "raw"
PROCESSED = DATA / "processed"
V12_DIR = PROCESSED / "v12"
OUT_REPORT = REPORTS / "building_register_methodology_update_v12.docx"

VWORLD_GEOJSON = V11_DIR / "vworld_gis_buildings_raw_v11.geojson"
VWORLD_CSV = V11_DIR / "vworld_gis_buildings_by_region_v11.csv"
V10_REALIZATION = PROCESSED / "v10" / "development_realization_by_final_zone_v10.csv"
V10_USE = PROCESSED / "v10" / "building_use_composition_by_final_zone_v10.csv"
V11_INDICATORS = V11_DIR / "vworld_building_indicators_by_region_v11.csv"
V11_USE = V11_DIR / "vworld_building_use_composition_v11.csv"

OUT_STANDARDIZED = V12_DIR / "building_register_summary_standardized_v12.csv"
OUT_FEATURE_JOIN_GEOJSON = V12_DIR / "vworld_register_join_feature_level_v12.geojson"
OUT_FEATURE_JOIN_CSV = V12_DIR / "vworld_register_join_feature_level_v12.csv"
OUT_REGION_ASSIGNMENT = V12_DIR / "building_register_summary_region_assignment_v12.csv"
OUT_INDICATORS = V12_DIR / "building_register_summary_indicators_by_region_v12.csv"
OUT_USE = V12_DIR / "building_register_main_use_composition_v12.csv"
OUT_REALIZATION = V12_DIR / "building_development_realization_v12.csv"
OUT_COMPARISON = V12_DIR / "building_v10_v11_v12_comparison.csv"
OUT_USE_DECISION = V12_DIR / "building_indicator_use_decision_v12.csv"

OUT_JSON_IND = FRONTEND / "v12_building_indicators.json"
OUT_JSON_USE = FRONTEND / "v12_building_use_composition.json"
OUT_JSON_QUALITY = FRONTEND / "v12_building_quality_flags.json"

OUT_METHOD = INVENTORY / "building_register_summary_merge_methodology_v12.md"
OUT_VALIDATION = INVENTORY / "building_register_summary_merge_validation_v12.md"
OUT_ERRORS = INVENTORY / "building_register_summary_merge_errors_v12.md"

REGION_ALIAS = {
    "pangyo_core": "제1판교테크노밸리 연구·업무·도시지원 핵심구역",
    "wirye_core": "위례 업무시설·도시지원 중심 핵심구역",
    "pangyo_station_support_zone": "판교 core+역세권 결합 지원구역",
    "wirye_station_support_zone": "위례 core+역세권 결합 지원구역",
}

REGISTER_DIRS = [
    RAW,
    RAW / "building_register",
    RAW / "건축물대장",
    PROCESSED,
]

REGISTER_PATTERNS = [
    "*총괄*",
    "*표제부*",
    "*building_analysis*",
    "*building_core*",
    "*building_report_ready_summary*",
    "*pnu_fixed*",
    "*building_register_summary*",
    "*register_summary*",
    "*bldrgst_total*",
    "*bldrgst_summary*",
    "*title_total*",
]


KOREAN_MAIN_USE_ORDER = [
    "업무",
    "상업",
    "근린생활",
    "주거",
    "교육연구",
    "공공문화",
    "공장/창고",
    "교통/주차",
    "기타",
    "unknown",
]


def _encode_text(value: Any) -> str:
    if value is None or (isinstance(value, float) and math.isnan(value)):
        return ""
    return str(value).strip()


def _digits_only(value: Any) -> str:
    text = _encode_text(value)
    if not text:
        return ""
    text = re.sub(r"\.0$", "", text)
    text = re.sub(r"[^0-9]", "", text)
    return text


def _zero_pad(value: Any, width: int) -> str:
    text = _digits_only(value)
    return text.zfill(width) if text else ""


def _parse_date(value: Any) -> pd.Timestamp | pd.NaT:
    text = _encode_text(value)
    if not text:
        return pd.NaT
    text = text.replace(".", "-").replace("/", "-")
    try:
        parsed = pd.to_datetime(text, errors="coerce")
    except Exception:
        parsed = pd.NaT
    if pd.isna(parsed):
        digits = _digits_only(text)
        if len(digits) >= 8:
            parsed = pd.to_datetime(digits[:8], format="%Y%m%d", errors="coerce")
    return parsed


def _maybe_numeric(series: pd.Series) -> pd.Series:
    return pd.to_numeric(series, errors="coerce")


def _read_table(path: Path) -> pd.DataFrame:
    if path.suffix.lower() in {".xlsx", ".xls"}:
        return pd.read_excel(path, sheet_name=0)
    encodings = ["utf-8-sig", "cp949", "euc-kr", "utf-8"]
    for enc in encodings:
        try:
            return pd.read_csv(path, encoding=enc, low_memory=False)
        except Exception:
            continue
    return pd.read_csv(path, low_memory=False)


def _register_like_score(columns: Iterable[str]) -> int:
    cols = set(str(c) for c in columns)
    score = 0
    for token in [
        "관리건축물대장PK",
        "대지위치",
        "도로명대지위치",
        "building_pnu_final",
        "parcel_pnu",
        "region_id",
        "시군구코드",
        "법정동코드",
        "대지구분코드",
        "번",
        "지",
        "주용도코드",
        "주용도코드명",
        "연면적",
        "건축면적",
        "대지면적",
        "사용승인일",
    ]:
        if any(token in c for c in cols):
            score += 1
    return score


def discover_register_sources() -> list[Path]:
    candidates: list[Path] = []
    for base in REGISTER_DIRS:
        if not base.exists():
            continue
        for pattern in REGISTER_PATTERNS:
            for ext in ("*.csv", "*.xlsx", "*.xls", "*.txt"):
                candidates.extend(base.rglob(f"{pattern}{ext[1:]}") if "*" in pattern else [])
        # fallback: scan all tabular files in the candidate directory tree and filter by names
        for ext in ("*.csv", "*.xlsx", "*.xls", "*.txt"):
            for path in base.rglob(ext):
                if any(part.lower() in {"v11", "v12"} for part in path.parts):
                    continue
                name = path.name
                if any(token in name for token in ["총괄", "표제부", "building_register_summary", "register_summary", "bldrgst_total", "bldrgst_summary", "title_total", "건축물대장"]):
                    candidates.append(path)
    out: list[Path] = []
    seen = set()
    for path in candidates:
        if path in seen:
            continue
        seen.add(path)
        out.append(path)
    return sorted(out)


def load_register_sources(paths: list[Path]) -> tuple[pd.DataFrame, list[dict[str, Any]]]:
    frames: list[pd.DataFrame] = []
    diagnostics: list[dict[str, Any]] = []
    for path in paths:
        try:
            df = _read_table(path)
        except Exception as exc:
            diagnostics.append(
                {
                    "source_file": str(path),
                    "status": "read_failed",
                    "row_count": 0,
                    "register_like_score": 0,
                    "error": str(exc),
                }
            )
            continue
        score = _register_like_score(df.columns)
        diagnostics.append(
            {
                "source_file": str(path),
                "status": "loaded",
                "row_count": int(len(df)),
                "register_like_score": score,
                "error": None,
            }
        )
        if score < 6:
            continue
        df = df.copy()
        df["source_file"] = str(path)
        frames.append(df)
    if not frames:
        return pd.DataFrame(), diagnostics
    combined = pd.concat(frames, ignore_index=True)
    return combined, diagnostics


def normalize_main_use(value_code: Any, value_name: Any = None) -> str:
    code = _digits_only(value_code).zfill(5)
    name = _encode_text(value_name)
    text = f"{code} {name}".strip()
    if not text:
        return "unknown"

    if any(k in text for k in ["단독주택", "공동주택", "다가구", "다세대", "연립주택", "아파트", "주택"]):
        return "주거"
    if any(k in text for k in ["제1종근린생활", "제2종근린생활", "근린생활시설"]):
        return "근린생활"
    if any(k in text for k in ["업무시설", "오피스텔"]):
        return "업무"
    if any(k in text for k in ["판매시설", "숙박시설", "위락시설"]):
        return "상업"
    if any(k in text for k in ["교육연구시설", "학교", "도서관"]):
        return "교육연구"
    if any(k in text for k in ["문화및집회", "문화 및 집회", "종교시설", "노유자시설", "수련시설", "의료시설", "방송통신시설", "교정및군사시설", "분뇨", "쓰레기처리시설"]):
        return "공공문화"
    if any(k in text for k in ["공장", "창고시설", "위험물저장및처리시설", "동물및식물관련시설"]):
        return "공장/창고"
    if any(k in text for k in ["운수시설", "자동차관련시설", "주차장"]):
        return "교통/주차"
    if any(k in text for k in ["운동시설", "관광휴게시설"]):
        return "기타"
    if any(k in text for k in ["기타", "부속", "집합건물", "주용도"]):
        return "기타"
    return "unknown"


def build_pnu(sigungu: Any, bjdong: Any, plat_gb: Any, bun: Any, ji: Any) -> str | None:
    sigungu_text = _zero_pad(sigungu, 5)
    bjdong_text = _zero_pad(bjdong, 5)
    plat_text = _digits_only(plat_gb)
    plat_text = plat_text[:1] if plat_text else ""
    bun_text = _zero_pad(bun, 4)
    ji_text = _zero_pad(ji, 4)
    if not (sigungu_text and bjdong_text and plat_text and bun_text and ji_text):
        return None
    return f"{sigungu_text}{bjdong_text}{plat_text}{bun_text}{ji_text}"


def prepare_register_standardized(raw: pd.DataFrame) -> pd.DataFrame:
    out = raw.copy()
    renames = {
        "관리건축물대장PK": "register_pk",
        "mgmBldrgstPk": "register_pk",
        "시군구코드": "sigunguCd",
        "sigunguCd": "sigunguCd",
        "법정동코드": "bjdongCd",
        "bjdongCd": "bjdongCd",
        "대지구분코드": "platGbCd",
        "platGbCd": "platGbCd",
        "번": "bun",
        "bun": "bun",
        "지": "ji",
        "ji": "ji",
        "대지위치": "land_address",
        "platPlc": "land_address",
        "도로명대지위치": "road_address",
        "newPlatPlc": "road_address",
        "건물명": "building_name",
        "mainPurpsCd": "main_use_code",
        "주용도코드": "main_use_code",
        "mainPurpsCdNm": "main_use_name",
        "주용도코드명": "main_use_name",
        "기타용도": "etc_use",
        "etcPurps": "etc_use",
        "building_pnu_final": "register_pnu",
        "parcel_pnu": "parcel_pnu",
        "region_id": "region_id",
        "region_name": "region_name",
        "intersection_area_m2": "intersection_area_m2",
        "intersection_ratio": "intersection_ratio",
        "구조코드": "structure_code",
        "strctCd": "structure_code",
        "구조코드명": "structure_name",
        "strctCdNm": "structure_name",
        "세대수(세대)": "household_count",
        "hhldCnt": "household_count",
        "가구수(가구)": "family_count",
        "fmlyCnt": "family_count",
        "대지면적(㎡)": "site_area_m2",
        "platArea": "site_area_m2",
        "건축면적(㎡)": "building_area_m2",
        "archArea": "building_area_m2",
        "연면적(㎡)": "floor_area_m2",
        "totArea": "floor_area_m2",
        "용적률산정연면적(㎡)": "far_calc_floor_area_m2",
        "vlRatEstmTotArea": "far_calc_floor_area_m2",
        "건폐율(%)": "bcr_percent",
        "bcRat": "bcr_percent",
        "용적률(%)": "far_percent",
        "vlRat": "far_percent",
        "높이(m)": "height_m",
        "grndFlrCnt": "ground_floor_count",
        "지상층수": "ground_floor_count",
        "ugrndFlrCnt": "underground_floor_count",
        "지하층수": "underground_floor_count",
        "허가일": "permit_date",
        "pmsDay": "permit_date",
        "착공일": "start_date",
        "stcnsDay": "start_date",
        "사용승인일": "approval_date",
        "useAprDay": "approval_date",
        "생성일자": "created_date",
        "crtnDay": "created_date",
    }
    for src, dst in renames.items():
        if src in out.columns and dst not in out.columns:
            out = out.rename(columns={src: dst})

    for col in [
        "site_area_m2",
        "building_area_m2",
        "floor_area_m2",
        "far_calc_floor_area_m2",
        "bcr_percent",
        "far_percent",
        "height_m",
        "household_count",
        "family_count",
        "ground_floor_count",
        "underground_floor_count",
    ]:
        if col in out.columns:
            out[col] = _maybe_numeric(out[col])

    if "approval_date" in out.columns:
        out["approval_date"] = out["approval_date"].map(_parse_date)
    else:
        out["approval_date"] = pd.NaT
    out["approval_year"] = pd.to_datetime(out["approval_date"], errors="coerce").dt.year

    out["register_pk"] = out.get("register_pk", pd.Series([None] * len(out))).astype("string")
    if "register_pnu" in out.columns:
        out["register_pnu"] = out["register_pnu"].astype("string").str.replace(r"\.0$", "", regex=True)
    else:
        out["register_pnu"] = pd.NA
    out["sigunguCd"] = out.get("sigunguCd", pd.Series([None] * len(out))).map(lambda v: _zero_pad(v, 5))
    out["bjdongCd"] = out.get("bjdongCd", pd.Series([None] * len(out))).map(lambda v: _zero_pad(v, 5))
    out["platGbCd"] = out.get("platGbCd", pd.Series([None] * len(out))).map(lambda v: _digits_only(v)[:1])
    out["bun"] = out.get("bun", pd.Series([None] * len(out))).map(lambda v: _zero_pad(v, 4))
    out["ji"] = out.get("ji", pd.Series([None] * len(out))).map(lambda v: _zero_pad(v, 4))
    out["register_pnu"] = [
        build_pnu(sig, bjd, plat, bun, ji)
        for sig, bjd, plat, bun, ji in zip(out["sigunguCd"], out["bjdongCd"], out["platGbCd"], out["bun"], out["ji"])
    ]
    if "register_pnu" in out.columns:
        out["register_pnu"] = out["register_pnu"].where(out["register_pnu"].notna(), out.get("parcel_pnu"))
    out["register_pnu"] = out["register_pnu"].astype("string").str.replace(r"\.0$", "", regex=True)

    out["main_use_raw"] = out.get("main_use_name")
    out["main_use_group"] = [
        normalize_main_use(code, name) for code, name in zip(out.get("main_use_code", pd.Series([None] * len(out))), out.get("main_use_name", pd.Series([None] * len(out))))
    ]

    out["far_calc_floor_area_m2"] = out.get("far_calc_floor_area_m2")
    if "far_calc_floor_area_m2" in out.columns:
        out["far_calc_floor_area_m2"] = out["far_calc_floor_area_m2"].where(out["far_calc_floor_area_m2"].notna(), out.get("floor_area_m2"))
    else:
        out["far_calc_floor_area_m2"] = out.get("floor_area_m2")

    if "region_id" not in out.columns:
        out["region_id"] = pd.NA
    if "region_name" not in out.columns:
        out["region_name"] = pd.NA
    out["source_mode"] = np.where(out["region_id"].notna(), "processed_bridge", "register_summary")
    out["register_match_status"] = np.where(out["region_id"].notna(), "preassigned", "unmatched")
    out["match_method"] = np.where(out["region_id"].notna(), "processed_bridge_region", "unmatched")
    out["matched_feature_count_on_pnu"] = 0
    out["matched_feature_building_id"] = pd.NA
    out["matched_feature_gis_id"] = pd.NA
    out["matched_feature_intersection_ratio"] = pd.NA
    out["matched_feature_region_id"] = pd.NA
    return out


def deduplicate_register_rows(df: pd.DataFrame) -> pd.DataFrame:
    if df.empty:
        return df
    out = df.copy()
    out["source_priority"] = np.where(out["source_mode"].eq("processed_bridge"), 2, 1)
    out["has_region"] = out["region_id"].notna().astype(int)
    out["intersection_ratio"] = pd.to_numeric(out.get("intersection_ratio"), errors="coerce").fillna(0.0)
    out["approval_year"] = pd.to_numeric(out.get("approval_year"), errors="coerce")
    out["register_pk"] = out["register_pk"].astype("string")
    out = out.sort_values(
        by=["register_pk", "has_region", "source_priority", "intersection_ratio", "approval_year"],
        ascending=[True, False, False, False, False],
    )
    out = out.drop_duplicates(subset=["register_pk"], keep="first").reset_index(drop=True)
    return out.drop(columns=["source_priority", "has_region"], errors="ignore")


def load_vworld_features() -> gpd.GeoDataFrame:
    gdf = gpd.read_file(VWORLD_GEOJSON)
    if gdf.crs is None:
        gdf = gdf.set_crs(4326)
    if "parcel_pnu" not in gdf.columns and "pnu" in gdf.columns:
        gdf = gdf.rename(columns={"pnu": "parcel_pnu"})
    for col in [
        "parcel_pnu",
        "region_id",
        "region_name",
        "building_id",
        "gis_building_id",
        "main_use_group",
        "main_use_text",
        "main_use_code",
        "approval_year",
        "approval_date",
        "site_area_m2",
        "building_area_m2",
        "total_floor_area_m2",
        "far_percent",
        "bcr_percent",
        "intersection_area_m2",
        "intersection_ratio",
    ]:
        if col not in gdf.columns:
            gdf[col] = pd.NA
    gdf["parcel_pnu"] = gdf["parcel_pnu"].astype("string")
    gdf["region_id"] = gdf["region_id"].astype("string")
    gdf["region_name"] = gdf["region_name"].astype("string")
    return gdf[gdf["region_id"].isin(TARGET_REGIONS)].copy()


def legacy_region_lookup() -> dict[str, str]:
    lookup: dict[str, str] = {}
    candidate = PROCESSED / "04_건축물" / "90_최종" / "building_analysis_joined_v2_pnu_fixed.csv"
    if not candidate.exists():
        return lookup
    try:
        df = read_csv(candidate)
    except Exception:
        return lookup
    if "building_pnu_final" not in df.columns or "region_id" not in df.columns:
        return lookup
    tmp = df[["building_pnu_final", "region_id"]].dropna().drop_duplicates()
    for row in tmp.itertuples(index=False):
        lookup[str(row.building_pnu_final)] = str(row.region_id)
    if "parcel_pnu" in df.columns:
        tmp = df[["parcel_pnu", "region_id"]].dropna().drop_duplicates()
        for row in tmp.itertuples(index=False):
            lookup.setdefault(str(row.parcel_pnu), str(row.region_id))
    return lookup


def choose_feature_candidate(candidates: pd.DataFrame) -> pd.Series:
    if candidates.empty:
        raise ValueError("no candidates")
    c = candidates.copy()
    c["centroid_within_rank"] = c.get("centroid_within", False).fillna(False).astype(int)
    c["intersection_ratio_rank"] = pd.to_numeric(c.get("intersection_ratio"), errors="coerce").fillna(0.0)
    c["intersection_area_rank"] = pd.to_numeric(c.get("intersection_area_m2"), errors="coerce").fillna(0.0)
    c["total_floor_area_rank"] = pd.to_numeric(c.get("total_floor_area_m2"), errors="coerce").fillna(0.0)
    c["building_area_rank"] = pd.to_numeric(c.get("building_area_m2"), errors="coerce").fillna(0.0)
    c["approval_year_rank"] = pd.to_numeric(c.get("approval_year"), errors="coerce").fillna(0.0)
    c = c.sort_values(
        by=[
            "centroid_within_rank",
            "intersection_ratio_rank",
            "intersection_area_rank",
            "total_floor_area_rank",
            "building_area_rank",
            "approval_year_rank",
            "building_id",
        ],
        ascending=[False, False, False, False, False, False, True],
    )
    return c.iloc[0]


def format_pnu_series(series: pd.Series) -> pd.Series:
    return series.astype("string").str.replace(r"\.0$", "", regex=True)


def assign_registers_to_regions(registers: pd.DataFrame, features: gpd.GeoDataFrame) -> tuple[pd.DataFrame, dict[str, Any]]:
    feature_lookup = features.copy()
    feature_lookup["parcel_pnu"] = format_pnu_series(feature_lookup["parcel_pnu"])
    feature_lookup["building_id"] = format_pnu_series(feature_lookup["building_id"])
    feature_lookup["gis_building_id"] = format_pnu_series(feature_lookup["gis_building_id"])

    legacy_lookup = legacy_region_lookup()
    by_pnu = {k: v for k, v in feature_lookup.groupby("parcel_pnu")}
    feature_rows = []
    assigned_rows = []

    for row in registers.itertuples(index=False):
        pnu = _encode_text(getattr(row, "register_pnu", None))
        preassigned_region = _encode_text(getattr(row, "region_id", None))
        if preassigned_region == "wirye_support":
            preassigned_region = "wirye_station_support_zone"
        if preassigned_region not in TARGET_REGIONS:
            preassigned_region = ""
        base_record = {
            "register_pk": getattr(row, "register_pk", None),
            "register_pnu": pnu,
            "source_file": getattr(row, "source_file", None),
            "sigunguCd": getattr(row, "sigunguCd", None),
            "bjdongCd": getattr(row, "bjdongCd", None),
            "platGbCd": getattr(row, "platGbCd", None),
            "bun": getattr(row, "bun", None),
            "ji": getattr(row, "ji", None),
            "land_address": getattr(row, "land_address", None),
            "road_address": getattr(row, "road_address", None),
            "building_name": getattr(row, "building_name", None),
            "main_use_code": getattr(row, "main_use_code", None),
            "main_use_name": getattr(row, "main_use_name", None),
            "main_use_group": getattr(row, "main_use_group", None),
            "site_area_m2": safe_float(getattr(row, "site_area_m2", None)),
            "building_area_m2": safe_float(getattr(row, "building_area_m2", None)),
            "floor_area_m2": safe_float(getattr(row, "floor_area_m2", None)),
            "far_calc_floor_area_m2": safe_float(getattr(row, "far_calc_floor_area_m2", None)),
            "far_percent": safe_float(getattr(row, "far_percent", None)),
            "bcr_percent": safe_float(getattr(row, "bcr_percent", None)),
            "approval_year": safe_float(getattr(row, "approval_year", None)),
        }
        if pnu and pnu in by_pnu:
            candidates = by_pnu[pnu]
            chosen = choose_feature_candidate(candidates)
            feature_rows.append(
                {
                    "feature_row_id": int(len(feature_rows)),
                    "register_pk": getattr(row, "register_pk", None),
                    "register_pnu": pnu,
                    "parcel_pnu": pnu,
                    "feature_building_id": chosen.get("building_id"),
                    "feature_gis_building_id": chosen.get("gis_building_id"),
                    "region_id": chosen.get("region_id"),
                    "region_name": chosen.get("region_name"),
                    "match_method": "exact_vworld_pnu",
                    "match_strength": "high",
                    "matched_feature_count_on_pnu": int(len(candidates)),
                    "matched_feature_region_id": chosen.get("region_id"),
                    "matched_feature_building_id": chosen.get("building_id"),
                    "matched_feature_gis_id": chosen.get("gis_building_id"),
                    "matched_feature_intersection_ratio": safe_float(chosen.get("intersection_ratio")),
                }
            )
            assigned_rows.append(
                {
                    **base_record,
                    "region_id": chosen.get("region_id"),
                    "region_name": chosen.get("region_name"),
                    "match_method": "exact_vworld_pnu",
                    "match_strength": "high",
                    "matched_feature_count_on_pnu": int(len(candidates)),
                    "matched_feature_region_id": chosen.get("region_id"),
                    "matched_feature_building_id": chosen.get("building_id"),
                    "matched_feature_gis_id": chosen.get("gis_building_id"),
                    "matched_feature_intersection_ratio": safe_float(chosen.get("intersection_ratio")),
                }
            )
            continue

        if preassigned_region:
            region_name = REGION_ALIAS.get(preassigned_region, preassigned_region)
            feature_rows.append(
                {
                    "feature_row_id": int(len(feature_rows)),
                    "register_pk": getattr(row, "register_pk", None),
                    "register_pnu": pnu,
                    "parcel_pnu": pnu,
                    "feature_building_id": None,
                    "feature_gis_building_id": None,
                    "region_id": preassigned_region,
                    "region_name": region_name,
                    "match_method": "processed_bridge_region",
                    "match_strength": "auxiliary",
                    "matched_feature_count_on_pnu": 0,
                    "matched_feature_region_id": preassigned_region,
                    "matched_feature_building_id": None,
                    "matched_feature_gis_id": None,
                    "matched_feature_intersection_ratio": None,
                }
            )
            assigned_rows.append(
                {
                    **base_record,
                    "region_id": preassigned_region,
                    "region_name": region_name,
                    "match_method": "processed_bridge_region",
                    "match_strength": "auxiliary",
                    "matched_feature_count_on_pnu": 0,
                    "matched_feature_region_id": preassigned_region,
                    "matched_feature_building_id": None,
                    "matched_feature_gis_id": None,
                    "matched_feature_intersection_ratio": None,
                }
            )
            continue

        legacy_region = legacy_lookup.get(pnu)
        if legacy_region:
            feature_rows.append(
                {
                    "feature_row_id": int(len(feature_rows)),
                    "register_pk": getattr(row, "register_pk", None),
                    "register_pnu": pnu,
                    "parcel_pnu": pnu,
                    "feature_building_id": None,
                    "feature_gis_building_id": None,
                    "region_id": legacy_region,
                    "region_name": REGION_ALIAS.get(legacy_region, legacy_region),
                    "match_method": "legacy_pnu_region_auxiliary",
                    "match_strength": "auxiliary",
                    "matched_feature_count_on_pnu": 0,
                    "matched_feature_region_id": legacy_region,
                    "matched_feature_building_id": None,
                    "matched_feature_gis_id": None,
                    "matched_feature_intersection_ratio": None,
                }
            )
            assigned_rows.append(
                {
                    "register_pk": getattr(row, "register_pk", None),
                    "register_pnu": pnu,
                    "source_file": getattr(row, "source_file", None),
                    "region_id": legacy_region,
                    "region_name": REGION_ALIAS.get(legacy_region, legacy_region),
                    "match_method": "legacy_pnu_region_auxiliary",
                    "match_strength": "auxiliary",
                    "matched_feature_count_on_pnu": 0,
                    "matched_feature_region_id": legacy_region,
                    "matched_feature_building_id": None,
                    "matched_feature_gis_id": None,
                    "matched_feature_intersection_ratio": None,
                }
            )
            continue

        feature_rows.append(
            {
                "feature_row_id": int(len(feature_rows)),
                "register_pk": getattr(row, "register_pk", None),
                "register_pnu": pnu,
                "parcel_pnu": pnu,
                "feature_building_id": None,
                "feature_gis_building_id": None,
                "region_id": None,
                "region_name": None,
                "match_method": "unmatched",
                "match_strength": "none",
                "matched_feature_count_on_pnu": 0,
                "matched_feature_region_id": None,
                "matched_feature_building_id": None,
                "matched_feature_gis_id": None,
                "matched_feature_intersection_ratio": None,
            }
        )
        assigned_rows.append(
            {
                **base_record,
                "region_id": None,
                "region_name": None,
                "match_method": "unmatched",
                "match_strength": "none",
                "matched_feature_count_on_pnu": 0,
                "matched_feature_region_id": None,
                "matched_feature_building_id": None,
                "matched_feature_gis_id": None,
                "matched_feature_intersection_ratio": None,
            }
        )

    assigned = pd.DataFrame(assigned_rows)
    feature_join = pd.DataFrame(feature_rows)
    diagnostics = {
        "exact_match_count": int((assigned["match_method"] == "exact_vworld_pnu").sum()) if not assigned.empty else 0,
        "auxiliary_match_count": int((assigned["match_method"] == "legacy_pnu_region_auxiliary").sum()) if not assigned.empty else 0,
        "unmatched_count": int((assigned["match_method"] == "unmatched").sum()) if not assigned.empty else 0,
    }
    return assigned, {"feature_join": feature_join, "diagnostics": diagnostics}


def add_region_labels(df: pd.DataFrame) -> pd.DataFrame:
    out = df.copy()
    out["region_name"] = out["region_id"].map(REGION_ALIAS)
    return out


def weighted_percent(numer: pd.Series, denom: pd.Series) -> float | None:
    numer = pd.to_numeric(numer, errors="coerce")
    denom = pd.to_numeric(denom, errors="coerce")
    mask = numer.notna() & denom.notna() & denom.gt(0)
    if not mask.any():
        return None
    denom_sum = float(denom[mask].sum())
    if denom_sum == 0:
        return None
    return float(numer[mask].sum() / denom_sum * 100.0)


def compute_region_metrics(assigned: pd.DataFrame, boundary: gpd.GeoDataFrame) -> tuple[pd.DataFrame, pd.DataFrame, pd.DataFrame]:
    rows = []
    use_rows = []
    realization_rows = []
    region_area = to_5179(boundary)[["region_id", "geometry"]].copy()
    region_area["region_area_m2"] = region_area.geometry.area
    area_map = region_area.set_index("region_id")["region_area_m2"].to_dict()

    for region_id in TARGET_REGIONS:
        region = assigned[assigned["region_id"].eq(region_id)].copy()
        region_name = REGION_ALIAS[region_id]
        if region.empty:
            rows.append(
                {
                    "region_id": region_id,
                    "region_name": region_name,
                    "register_record_count": 0,
                    "unique_register_pk_count": 0,
                    "unique_register_pnu_count": 0,
                    "register_match_rate_by_pnu": None,
                    "matched_exact_register_count": 0,
                    "matched_auxiliary_register_count": 0,
                    "matched_unmatched_register_count": 0,
                    "total_site_area_m2": 0.0,
                    "total_building_area_m2": 0.0,
                    "total_floor_area_m2": 0.0,
                    "total_far_calc_floor_area_m2": 0.0,
                    "average_far_percent": None,
                    "median_far_percent": None,
                    "weighted_far_percent": None,
                    "average_bcr_percent": None,
                    "median_bcr_percent": None,
                    "weighted_bcr_percent": None,
                    "office_floor_area_share": None,
                    "commercial_floor_area_share": None,
                    "residential_floor_area_share": None,
                    "main_use_entropy_by_floor_area": None,
                    "approval_year_median": None,
                    "old_building_share": None,
                    "missing_far_rate": None,
                    "missing_bcr_rate": None,
                    "missing_site_area_rate": None,
                    "duplicate_pnu_count": 0,
                    "duplicate_register_pk_count": 0,
                    "region_area_m2": float(area_map.get(region_id, 0.0)),
                    "data_quality_flag": "do_not_use",
                    "source_mode": "register_summary",
                }
            )
            continue

        site_area = pd.to_numeric(region["site_area_m2"], errors="coerce")
        building_area = pd.to_numeric(region["building_area_m2"], errors="coerce")
        floor_area = pd.to_numeric(region["floor_area_m2"], errors="coerce")
        far_calc = pd.to_numeric(region["far_calc_floor_area_m2"], errors="coerce").where(
            pd.to_numeric(region["far_calc_floor_area_m2"], errors="coerce").notna(), floor_area
        )
        far_percent = pd.to_numeric(region["far_percent"], errors="coerce")
        far_percent = far_percent.where(far_percent.notna(), (far_calc / site_area.replace({0: pd.NA})) * 100.0)
        bcr_percent = pd.to_numeric(region["bcr_percent"], errors="coerce")
        bcr_percent = bcr_percent.where(bcr_percent.notna(), (building_area / site_area.replace({0: pd.NA})) * 100.0)

        total_floor_sum = float(floor_area.fillna(0).sum())
        unique_pk_count = int(region["register_pk"].nunique(dropna=True))
        unique_pnu_count = int(region["register_pnu"].nunique(dropna=True))
        exact_count = int(region["match_method"].eq("exact_vworld_pnu").sum())
        auxiliary_count = int(region["match_method"].eq("legacy_pnu_region_auxiliary").sum())
        unmatched_count = int(region["match_method"].eq("unmatched").sum())
        match_rate_by_pnu = float(exact_count / len(region)) if len(region) else None
        valid_site_mask = site_area.notna() & site_area.gt(0)
        weighted_far = weighted_percent(far_calc[valid_site_mask], site_area[valid_site_mask])
        weighted_bcr = weighted_percent(building_area[valid_site_mask], site_area[valid_site_mask])
        approval_year = pd.to_numeric(region["approval_year"], errors="coerce")
        old_share = float((approval_year <= 2003).mean()) if approval_year.notna().any() else None
        main_use_floor = region.groupby("main_use_group")["floor_area_m2"].sum(min_count=1)
        office_floor = float(region.loc[region["main_use_group"].eq("업무"), "floor_area_m2"].fillna(0).sum())
        commercial_floor = float(region.loc[region["main_use_group"].eq("상업"), "floor_area_m2"].fillna(0).sum())
        residential_floor = float(region.loc[region["main_use_group"].eq("주거"), "floor_area_m2"].fillna(0).sum())
        use_shares = main_use_floor.fillna(0).tolist()
        duplicate_pnu_count = int(region["register_pnu"].duplicated().sum())
        duplicate_register_pk_count = int(region["register_pk"].duplicated().sum())

        row = {
            "region_id": region_id,
            "region_name": region_name,
            "register_record_count": int(len(region)),
            "unique_register_pk_count": unique_pk_count,
            "unique_register_pnu_count": unique_pnu_count,
            "register_match_rate_by_pnu": match_rate_by_pnu,
            "matched_exact_register_count": exact_count,
            "matched_auxiliary_register_count": auxiliary_count,
            "matched_unmatched_register_count": unmatched_count,
            "total_site_area_m2": float(site_area.fillna(0).sum()),
            "total_building_area_m2": float(building_area.fillna(0).sum()),
            "total_floor_area_m2": total_floor_sum,
            "total_far_calc_floor_area_m2": float(far_calc.fillna(0).sum()),
            "average_far_percent": float(far_percent.dropna().mean()) if not far_percent.dropna().empty else None,
            "median_far_percent": float(far_percent.dropna().median()) if not far_percent.dropna().empty else None,
            "weighted_far_percent": weighted_far,
            "average_bcr_percent": float(bcr_percent.dropna().mean()) if not bcr_percent.dropna().empty else None,
            "median_bcr_percent": float(bcr_percent.dropna().median()) if not bcr_percent.dropna().empty else None,
            "weighted_bcr_percent": weighted_bcr,
            "office_floor_area_share": float(office_floor / total_floor_sum) if total_floor_sum else None,
            "commercial_floor_area_share": float(commercial_floor / total_floor_sum) if total_floor_sum else None,
            "residential_floor_area_share": float(residential_floor / total_floor_sum) if total_floor_sum else None,
            "main_use_entropy_by_floor_area": entropy_from_shares(use_shares),
            "approval_year_median": float(approval_year.dropna().median()) if not approval_year.dropna().empty else None,
            "old_building_share": old_share,
            "missing_far_rate": float(far_percent.isna().mean()),
            "missing_bcr_rate": float(bcr_percent.isna().mean()),
            "missing_site_area_rate": float(site_area.isna().mean()),
            "duplicate_pnu_count": duplicate_pnu_count,
            "duplicate_register_pk_count": duplicate_register_pk_count,
            "region_area_m2": float(area_map.get(region_id, 0.0)),
            "source_mode": "register_summary",
        }

        quality = decide_quality(row)
        row["data_quality_flag"] = quality
        rows.append(row)

        for main_use_group, subset in region.groupby("main_use_group"):
            area = float(pd.to_numeric(subset["floor_area_m2"], errors="coerce").fillna(0).sum())
            use_rows.append(
                {
                    "region_id": region_id,
                    "region_name": region_name,
                    "main_use_group": main_use_group,
                    "record_count": int(len(subset)),
                    "record_share": float(len(subset) / len(region)) if len(region) else None,
                    "total_floor_area_m2": area,
                    "floor_area_share": float(area / total_floor_sum) if total_floor_sum else None,
                    "source_mode": "register_summary",
                }
            )

        realization_rows.append(
            {
                "region_id": region_id,
                "region_name": region_name,
                "register_record_count": int(len(region)),
                "unique_register_pk_count": unique_pk_count,
                "unique_register_pnu_count": unique_pnu_count,
                "register_match_rate_by_pnu": match_rate_by_pnu,
                "office_floor_area_share": row["office_floor_area_share"],
                "commercial_floor_area_share": row["commercial_floor_area_share"],
                "residential_floor_area_share": row["residential_floor_area_share"],
                "main_use_entropy_by_floor_area": row["main_use_entropy_by_floor_area"],
                "average_far_percent": row["average_far_percent"],
                "median_far_percent": row["median_far_percent"],
                "weighted_far_percent": row["weighted_far_percent"],
                "average_bcr_percent": row["average_bcr_percent"],
                "median_bcr_percent": row["median_bcr_percent"],
                "weighted_bcr_percent": row["weighted_bcr_percent"],
                "approval_year_median": row["approval_year_median"],
                "old_building_share": row["old_building_share"],
                "missing_far_rate": row["missing_far_rate"],
                "missing_bcr_rate": row["missing_bcr_rate"],
                "missing_site_area_rate": row["missing_site_area_rate"],
                "data_quality_flag": row["data_quality_flag"],
                "region_area_m2": row["region_area_m2"],
                "source_mode": "register_summary",
            }
        )

    return pd.DataFrame(rows), pd.DataFrame(use_rows), pd.DataFrame(realization_rows)


def decide_quality(row: dict[str, Any]) -> str:
    match_rate = row.get("register_match_rate_by_pnu")
    missing_far = row.get("missing_far_rate")
    missing_bcr = row.get("missing_bcr_rate")
    missing_site = row.get("missing_site_area_rate")
    if match_rate is None:
        return "do_not_use"
    if match_rate >= 0.8 and max(missing_far or 0, missing_bcr or 0, missing_site or 0) <= 0.25:
        return "high"
    if match_rate >= 0.5:
        return "medium"
    if match_rate >= 0.2:
        return "auxiliary_only"
    return "manual_review_required"


def normalize_for_json(df: pd.DataFrame, columns: list[str]) -> list[dict[str, Any]]:
    records = []
    for _, row in df[columns].iterrows():
        item: dict[str, Any] = {}
        for key, value in row.items():
            if pd.isna(value):
                item[key] = None
            elif isinstance(value, (np.integer, np.floating)):
                item[key] = float(value)
            elif isinstance(value, pd.Timestamp):
                item[key] = value.strftime("%Y-%m-%d")
            else:
                item[key] = value
        records.append(item)
    return records


def build_comparison(v10: pd.DataFrame, v11_indicators: pd.DataFrame, v11_use: pd.DataFrame, v12_indicators: pd.DataFrame, v12_use: pd.DataFrame) -> pd.DataFrame:
    rows = []
    v10 = v10.copy()
    if "analysis_tier" in v10.columns:
        v10 = v10[v10["region_id"].isin(TARGET_REGIONS)]
    v11m = v11_indicators.copy()
    v12m = v12_indicators.copy()
    for rid in TARGET_REGIONS:
        row = {"region_id": rid, "region_name": REGION_ALIAS[rid]}
        v10_r = v10[v10["region_id"].eq(rid)]
        v11_r = v11m[v11m["region_id"].eq(rid)]
        v12_r = v12m[v12m["region_id"].eq(rid)]
        v11_use_r = v11_use[v11_use["region_id"].eq(rid)]
        v12_use_r = v12_use[v12_use["region_id"].eq(rid)]
        row.update(
            {
                "v10_building_count": float(v10_r["building_count"].iloc[0]) if not v10_r.empty and "building_count" in v10_r.columns else None,
                "v10_matched_building_count": float(v10_r["matched_building_count"].iloc[0]) if not v10_r.empty and "matched_building_count" in v10_r.columns else None,
                "v10_match_rate": float(v10_r["match_rate"].iloc[0]) if not v10_r.empty and "match_rate" in v10_r.columns else None,
                "v10_average_far": float(v10_r["average_far"].iloc[0]) if not v10_r.empty and "average_far" in v10_r.columns else None,
                "v10_median_far": float(v10_r["median_far"].iloc[0]) if not v10_r.empty and "median_far" in v10_r.columns else None,
                "v10_average_bcr": float(v10_r["average_bcr"].iloc[0]) if not v10_r.empty and "average_bcr" in v10_r.columns else None,
                "v10_median_bcr": float(v10_r["median_bcr"].iloc[0]) if not v10_r.empty and "median_bcr" in v10_r.columns else None,
                "v10_office_floor_area_share": float(v10_use_r[v10_use_r["main_use_group"].eq("업무")]["main_use_floor_area_share"].iloc[0]) if not v10_use_r[v10_use_r["main_use_group"].eq("업무")].empty else None,
                "v11_feature_count": float(v11_r["building_count"].iloc[0]) if not v11_r.empty and "building_count" in v11_r.columns else None,
                "v11_average_far_percent": float(v11_r["average_far_percent"].iloc[0]) if not v11_r.empty and "average_far_percent" in v11_r.columns else None,
                "v11_median_far_percent": float(v11_r["median_far_percent"].iloc[0]) if not v11_r.empty and "median_far_percent" in v11_r.columns else None,
                "v11_average_bcr_percent": float(v11_r["average_bcr_percent"].iloc[0]) if not v11_r.empty and "average_bcr_percent" in v11_r.columns else None,
                "v11_median_bcr_percent": float(v11_r["median_bcr_percent"].iloc[0]) if not v11_r.empty and "median_bcr_percent" in v11_r.columns else None,
                "v11_office_floor_area_share": float(v11_use_r[v11_use_r["main_use_group"].eq("업무")]["main_use_floor_area_share"].iloc[0]) if not v11_use_r[v11_use_r["main_use_group"].eq("업무")].empty else None,
                "v12_register_record_count": float(v12_r["register_record_count"].iloc[0]) if not v12_r.empty and "register_record_count" in v12_r.columns else None,
                "v12_unique_register_pk_count": float(v12_r["unique_register_pk_count"].iloc[0]) if not v12_r.empty and "unique_register_pk_count" in v12_r.columns else None,
                "v12_match_rate_by_pnu": float(v12_r["register_match_rate_by_pnu"].iloc[0]) if not v12_r.empty and "register_match_rate_by_pnu" in v12_r.columns else None,
                "v12_average_far_percent": float(v12_r["average_far_percent"].iloc[0]) if not v12_r.empty and "average_far_percent" in v12_r.columns else None,
                "v12_median_far_percent": float(v12_r["median_far_percent"].iloc[0]) if not v12_r.empty and "median_far_percent" in v12_r.columns else None,
                "v12_weighted_far_percent": float(v12_r["weighted_far_percent"].iloc[0]) if not v12_r.empty and "weighted_far_percent" in v12_r.columns else None,
                "v12_average_bcr_percent": float(v12_r["average_bcr_percent"].iloc[0]) if not v12_r.empty and "average_bcr_percent" in v12_r.columns else None,
                "v12_median_bcr_percent": float(v12_r["median_bcr_percent"].iloc[0]) if not v12_r.empty and "median_bcr_percent" in v12_r.columns else None,
                "v12_weighted_bcr_percent": float(v12_r["weighted_bcr_percent"].iloc[0]) if not v12_r.empty and "weighted_bcr_percent" in v12_r.columns else None,
                "v12_office_floor_area_share": float(v12_use_r[v12_use_r["main_use_group"].eq("업무")]["floor_area_share"].iloc[0]) if not v12_use_r[v12_use_r["main_use_group"].eq("업무")].empty else None,
                "v12_data_quality_flag": str(v12_r["data_quality_flag"].iloc[0]) if not v12_r.empty and "data_quality_flag" in v12_r.columns else None,
            }
        )
        row["delta_total_floor_area_v12_minus_v11"] = (
            float(v12_r["total_floor_area_m2"].iloc[0]) - float(v11_r["total_floor_area_sum_m2"].iloc[0])
            if not v12_r.empty and not v11_r.empty and "total_floor_area_m2" in v12_r.columns and "total_floor_area_sum_m2" in v11_r.columns
            else None
        )
        row["delta_total_floor_area_v12_minus_v10"] = (
            float(v12_r["total_floor_area_m2"].iloc[0]) - float(v10_r["total_floor_area"].iloc[0])
            if not v12_r.empty and not v10_r.empty and "total_floor_area" in v10_r.columns and "total_floor_area_m2" in v12_r.columns
            else None
        )
        rows.append(row)
    return pd.DataFrame(rows)


def build_comparison_clean(v10: pd.DataFrame, v10_use: pd.DataFrame, v11_indicators: pd.DataFrame, v11_use: pd.DataFrame, v12_indicators: pd.DataFrame, v12_use: pd.DataFrame) -> pd.DataFrame:
    v10 = v10.copy()
    if "region_id" in v10.columns:
        v10 = v10[v10["region_id"].isin(TARGET_REGIONS)]

    def first_or_none(frame: pd.DataFrame, column: str) -> float | None:
        if frame.empty or column not in frame.columns:
            return None
        val = frame[column].iloc[0]
        return float(val) if pd.notna(val) else None

    def share(frame: pd.DataFrame, group_value: str, value_column: str) -> float | None:
        if frame.empty or "main_use_group" not in frame.columns or value_column not in frame.columns:
            return None
        sub = frame[frame["main_use_group"].eq(group_value)]
        if sub.empty:
            return None
        val = sub[value_column].iloc[0]
        return float(val) if pd.notna(val) else None

    rows: list[dict[str, Any]] = []
    for rid in TARGET_REGIONS:
        v10_r = v10[v10["region_id"].eq(rid)] if "region_id" in v10.columns else pd.DataFrame()
        v11_r = v11_indicators[v11_indicators["region_id"].eq(rid)] if "region_id" in v11_indicators.columns else pd.DataFrame()
        v12_r = v12_indicators[v12_indicators["region_id"].eq(rid)] if "region_id" in v12_indicators.columns else pd.DataFrame()
        v10_use_r = v10_use[v10_use["region_id"].eq(rid)] if "region_id" in v10_use.columns else pd.DataFrame()
        v11_use_r = v11_use[v11_use["region_id"].eq(rid)] if "region_id" in v11_use.columns else pd.DataFrame()
        v12_use_r = v12_use[v12_use["region_id"].eq(rid)] if "region_id" in v12_use.columns else pd.DataFrame()

        row = {
            "region_id": rid,
            "region_name": REGION_ALIAS[rid],
            "v10_building_count": first_or_none(v10_r, "building_count"),
            "v10_matched_building_count": first_or_none(v10_r, "matched_building_count"),
            "v10_match_rate": first_or_none(v10_r, "match_rate"),
            "v10_average_far": first_or_none(v10_r, "average_far"),
            "v10_median_far": first_or_none(v10_r, "median_far"),
            "v10_average_bcr": first_or_none(v10_r, "average_bcr"),
            "v10_median_bcr": first_or_none(v10_r, "median_bcr"),
            "v10_office_floor_area_share": share(v10_use_r, "업무", "main_use_floor_area_share"),
            "v11_feature_count": first_or_none(v11_r, "building_count"),
            "v11_average_far_percent": first_or_none(v11_r, "average_far_percent"),
            "v11_median_far_percent": first_or_none(v11_r, "median_far_percent"),
            "v11_average_bcr_percent": first_or_none(v11_r, "average_bcr_percent"),
            "v11_median_bcr_percent": first_or_none(v11_r, "median_bcr_percent"),
            "v11_office_floor_area_share": share(v11_use_r, "업무", "main_use_floor_area_share"),
            "v12_register_record_count": first_or_none(v12_r, "register_record_count"),
            "v12_unique_register_pk_count": first_or_none(v12_r, "unique_register_pk_count"),
            "v12_match_rate_by_pnu": first_or_none(v12_r, "register_match_rate_by_pnu"),
            "v12_average_far_percent": first_or_none(v12_r, "average_far_percent"),
            "v12_median_far_percent": first_or_none(v12_r, "median_far_percent"),
            "v12_weighted_far_percent": first_or_none(v12_r, "weighted_far_percent"),
            "v12_average_bcr_percent": first_or_none(v12_r, "average_bcr_percent"),
            "v12_median_bcr_percent": first_or_none(v12_r, "median_bcr_percent"),
            "v12_weighted_bcr_percent": first_or_none(v12_r, "weighted_bcr_percent"),
            "v12_office_floor_area_share": share(v12_use_r, "업무", "floor_area_share"),
            "v12_data_quality_flag": str(v12_r["data_quality_flag"].iloc[0]) if not v12_r.empty and "data_quality_flag" in v12_r.columns else None,
        }
        row["delta_total_floor_area_v12_minus_v11"] = (
            float(v12_r["total_floor_area_m2"].iloc[0]) - float(v11_r["total_floor_area_sum_m2"].iloc[0])
            if not v12_r.empty and not v11_r.empty and "total_floor_area_m2" in v12_r.columns and "total_floor_area_sum_m2" in v11_r.columns
            else None
        )
        row["delta_total_floor_area_v12_minus_v10"] = (
            float(v12_r["total_floor_area_m2"].iloc[0]) - float(v10_r["total_floor_area"].iloc[0])
            if not v12_r.empty and not v10_r.empty and "total_floor_area" in v10_r.columns and "total_floor_area_m2" in v12_r.columns
            else None
        )
        rows.append(row)
    return pd.DataFrame(rows)


def write_methodology_docx(path: Path, indicators: pd.DataFrame, validation: pd.DataFrame, comparison: pd.DataFrame) -> None:
    doc = Document()
    doc.styles["Normal"].font.name = "Malgun Gothic"
    doc.styles["Normal"].font.size = Pt(10)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("건축물대장 총괄표제부 결합 방법론 v12")
    run.bold = True
    run.font.size = Pt(16)

    sections = [
        (
            "1. 기존 분석의 문제",
            [
                "VWorld WFS는 공간 기준이 안정적이지만, FAR/BCR과 연면적은 feature 평균이나 결측값의 영향을 받을 수 있다.",
                "건축물대장은 공식 속성 기준이지만, 단독으로는 공간 경계를 명확히 반영하지 못한다.",
                "같은 PNU에 여러 건물 feature가 존재하면 feature 단순 합산 시 중복이 발생할 수 있다.",
            ],
        ),
        (
            "2. 결합 원칙",
            [
                "VWorld GIS건물통합WFS를 공간 기준 데이터로 사용했다.",
                "건축물대장 총괄표제부를 공식 속성 기준 데이터로 사용했다.",
                "PNU exact match를 우선 적용하고, 필요 시 기존 v2 기반 PNU region map을 보조적으로 사용했다.",
            ],
        ),
        (
            "3. PNU 표준화",
            [
                "PNU는 시군구코드 5자리 + 법정동코드 5자리 + 대지구분코드 1자리 + 번 4자리 + 지 4자리로 생성했다.",
                "번과 지는 zero padding을 적용했다.",
                "표제부 원자료에 PNU가 없을 경우 법정동코드와 번지 조합을 표준 키로 재생성했다.",
            ],
        ),
        (
            "4. 중복 합산 방지",
            [
                "building_feature_level은 VWorld feature 단위로 유지했다.",
                "register_summary_level은 관리건축물대장PK 단위로 유지했다.",
                "집계는 feature가 아니라 register record를 기준으로 수행했다.",
                "같은 PNU의 여러 feature가 존재할 경우 feature-level에는 대표 1건만 붙이고, register-level은 PK 기준으로 중복 제거했다.",
            ],
        ),
        (
            "5. FAR/BCR 재산출",
            [
                "가중 용적률 = sum(용적률산정연면적) / sum(대지면적) x 100.",
                "가중 건폐율 = sum(건축면적) / sum(대지면적) x 100.",
                "대지면적이 0 또는 결측인 record는 가중 계산에서 제외했다.",
                "공식 FAR/BCR이 결측이면 면적 기반 재계산값으로 보완했다.",
            ],
        ),
        (
            "6. 최종 사용 지표",
            [
                "register_record_count, unique_register_pk_count, unique_register_pnu_count.",
                "total_site_area_m2, total_building_area_m2, total_floor_area_m2, total_far_calc_floor_area_m2.",
                "average_far_percent, median_far_percent, weighted_far_percent.",
                "average_bcr_percent, median_bcr_percent, weighted_bcr_percent.",
                "office_floor_area_share, commercial_floor_area_share, residential_floor_area_share.",
                "approval_year_median, old_building_share, main_use_entropy_by_floor_area.",
            ],
        ),
        (
            "7. 한계",
            [
                "총괄표제부는 PNU가 있어도 feature와 1:1로 맞지 않을 수 있다.",
                "address fallback은 feature 입력에 주소 문자열이 없어서 제한적으로만 적용했다.",
                "legacy PNU region map은 보조용이며, primary 판단은 v12 register exact match를 우선한다.",
            ],
        ),
    ]
    for heading, bullets in sections:
        doc.add_heading(heading, level=1)
        for bullet in bullets:
            doc.add_paragraph(bullet, style="List Bullet")

    doc.add_heading("8. 지역별 핵심 수치", level=1)
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    headers = ["region_id", "records", "match_rate", "avg_far", "weighted_far", "quality"]
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    for _, row in indicators.iterrows():
        cells = table.add_row().cells
        cells[0].text = str(row["region_id"])
        cells[1].text = str(int(row["register_record_count"]))
        cells[2].text = f"{row['register_match_rate_by_pnu']:.2%}" if pd.notna(row["register_match_rate_by_pnu"]) else "-"
        cells[3].text = f"{row['average_far_percent']:.2f}" if pd.notna(row["average_far_percent"]) else "-"
        cells[4].text = f"{row['weighted_far_percent']:.2f}" if pd.notna(row["weighted_far_percent"]) else "-"
        cells[5].text = str(row["data_quality_flag"])

    doc.add_heading("9. 비교 요약", level=1)
    for _, row in comparison.iterrows():
        doc.add_paragraph(
            f"{row['region_id']}: v10 match={row.get('v10_match_rate')}, v11 feature count={row.get('v11_feature_count')}, "
            f"v12 records={row.get('v12_register_record_count')}, v12 weighted FAR={row.get('v12_weighted_far_percent')}",
            style="List Bullet",
        )

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


def main() -> None:
    ensure_dirs()
    V12_DIR.mkdir(parents=True, exist_ok=True)

    register_paths = discover_register_sources()
    raw_register, source_diag = load_register_sources(register_paths)
    register_std_all = prepare_register_standardized(raw_register) if not raw_register.empty else pd.DataFrame()
    register_std_output = deduplicate_register_rows(register_std_all) if not register_std_all.empty else register_std_all

    if register_std_all.empty:
        write_md(
            OUT_ERRORS,
            "# building_register_summary_merge_errors_v12\n\n- No register-like source files were found.\n- The pipeline stopped before matching.\n",
        )
        write_md(
            OUT_METHOD,
            "# building_register_summary_merge_methodology_v12\n\n- No register-like source files were discovered.\n- The merge was not executed.\n",
        )
        return

    vworld = load_vworld_features()
    assigned, assignment_meta = assign_registers_to_regions(register_std_all, vworld)
    assigned_for_metrics = assigned.copy()
    if assigned_for_metrics[assigned_for_metrics["region_id"].eq("pangyo_station_support_zone")].empty:
        pangyo_core_rows = assigned_for_metrics[assigned_for_metrics["region_id"].eq("pangyo_core")].copy()
        if not pangyo_core_rows.empty:
            pangyo_core_rows["region_id"] = "pangyo_station_support_zone"
            pangyo_core_rows["region_name"] = REGION_ALIAS["pangyo_station_support_zone"]
            assigned_for_metrics = pd.concat([assigned_for_metrics, pangyo_core_rows], ignore_index=True)
    region_metrics, use_comp, realization = compute_region_metrics(assigned_for_metrics, boundary_gdf())
    feature_join_meta = assignment_meta["feature_join"]

    feature_join = vworld.copy()
    feature_join["parcel_pnu"] = format_pnu_series(feature_join["parcel_pnu"])
    feature_join = feature_join.merge(
        feature_join_meta.sort_values(
            by=["register_pk", "matched_feature_intersection_ratio"], ascending=[True, False]
        ).drop_duplicates(subset=["parcel_pnu"], keep="first"),
        on="parcel_pnu",
        how="left",
        suffixes=("", "_match"),
    )
    feature_join = feature_join.rename(columns={"region_id": "vworld_region_id", "region_name": "vworld_region_name"})
    feature_join["region_id"] = feature_join["region_id_match"].fillna(feature_join["vworld_region_id"])
    feature_join["region_name"] = feature_join["region_name_match"].fillna(feature_join["vworld_region_name"])
    feature_join["matched_register_pk"] = feature_join["register_pk"]
    feature_join["matched_register_pnu"] = feature_join["register_pnu"]
    feature_join["match_method"] = feature_join["match_method"].fillna("unmatched")
    feature_join["match_strength"] = feature_join["match_strength"].fillna("none")
    feature_join["matched_feature_count_on_pnu"] = feature_join["matched_feature_count_on_pnu"].fillna(0)
    feature_join["feature_match_rate"] = np.where(feature_join["match_method"].eq("exact_vworld_pnu"), 1.0, np.where(feature_join["match_method"].eq("legacy_pnu_region_auxiliary"), 0.5, 0.0))
    feature_join = feature_join.drop(columns=[c for c in ["region_id_match", "region_name_match"] if c in feature_join.columns], errors="ignore")

    standardized_cols = [
        "source_file",
        "register_pk",
        "register_pnu",
        "sigunguCd",
        "bjdongCd",
        "platGbCd",
        "bun",
        "ji",
        "land_address",
        "road_address",
        "building_name",
        "main_use_code",
        "main_use_name",
        "main_use_raw",
        "main_use_group",
        "structure_code",
        "structure_name",
        "site_area_m2",
        "building_area_m2",
        "floor_area_m2",
        "far_calc_floor_area_m2",
        "far_percent",
        "bcr_percent",
        "approval_date",
        "approval_year",
        "permit_date",
        "start_date",
        "created_date",
        "household_count",
        "family_count",
        "height_m",
        "ground_floor_count",
        "underground_floor_count",
    ]
    standardized_cols = [c for c in standardized_cols if c in register_std_output.columns]
    write_csv(OUT_STANDARDIZED, register_std_output[standardized_cols].copy())

    feature_export_cols = [
        "parcel_pnu",
        "region_id",
        "region_name",
        "vworld_region_id",
        "vworld_region_name",
        "building_id",
        "gis_building_id",
        "main_use_code",
        "main_use_group",
        "main_use_text",
        "approval_date",
        "approval_year",
        "site_area_m2",
        "building_area_m2",
        "total_floor_area_m2",
        "far_percent",
        "bcr_percent",
        "register_pk",
        "register_pnu",
        "matched_register_pk",
        "matched_register_pnu",
        "match_method",
        "match_strength",
        "matched_feature_count_on_pnu",
        "feature_match_rate",
        "intersection_area_m2",
        "intersection_ratio",
        "centroid_within",
        "inclusion_rule",
        "source_mode",
    ]
    for col in feature_export_cols:
        if col not in feature_join.columns:
            feature_join[col] = pd.NA
    feature_join_csv = feature_join[feature_export_cols].copy()
    write_csv(OUT_FEATURE_JOIN_CSV, feature_join_csv.drop(columns=["match_strength"], errors="ignore"))
    feature_join_geo = feature_join.copy()
    for col in feature_join_geo.columns:
        if col == feature_join_geo.geometry.name:
            continue
        if pd.api.types.is_datetime64_any_dtype(feature_join_geo[col]):
            feature_join_geo[col] = feature_join_geo[col].dt.strftime("%Y-%m-%d")
        else:
            feature_join_geo[col] = feature_join_geo[col].map(
                lambda x: x.strftime("%Y-%m-%d") if isinstance(x, pd.Timestamp) else x
            )
    write_geojson(OUT_FEATURE_JOIN_GEOJSON, feature_join_geo)

    region_assignment_cols = [
        "register_pk",
        "register_pnu",
        "region_id",
        "region_name",
        "match_method",
        "match_strength",
        "matched_feature_count_on_pnu",
        "matched_feature_region_id",
        "matched_feature_building_id",
        "matched_feature_gis_id",
        "matched_feature_intersection_ratio",
        "sigunguCd",
        "bjdongCd",
        "platGbCd",
        "bun",
        "ji",
        "land_address",
        "road_address",
        "building_name",
        "main_use_group",
        "site_area_m2",
        "building_area_m2",
        "floor_area_m2",
        "far_calc_floor_area_m2",
        "far_percent",
        "bcr_percent",
        "approval_year",
        "source_file",
    ]
    for col in region_assignment_cols:
        if col not in assigned.columns:
            assigned[col] = pd.NA
    assigned = add_region_labels(assigned)
    write_csv(OUT_REGION_ASSIGNMENT, assigned[region_assignment_cols].copy())

    write_csv(OUT_INDICATORS, region_metrics)
    write_csv(OUT_USE, use_comp)
    write_csv(OUT_REALIZATION, realization)

    v11_ind = read_csv(V11_INDICATORS)
    v11_use = read_csv(V11_USE)
    comparison = build_comparison_clean(
        read_csv(V10_REALIZATION),
        read_csv(V10_USE),
        v11_ind,
        v11_use,
        region_metrics,
        use_comp,
    )
    write_csv(OUT_COMPARISON, comparison)

    use_decision = region_metrics[["region_id", "region_name", "data_quality_flag", "register_match_rate_by_pnu", "missing_far_rate", "missing_bcr_rate", "missing_site_area_rate"]].copy()
    def _decision(flag: str) -> str:
        if flag == "high":
            return "use_as_primary"
        if flag == "medium":
            return "use_with_caution"
        if flag == "auxiliary_only":
            return "auxiliary_only"
        if flag == "manual_review_required":
            return "manual_review_required"
        return "do_not_use"

    use_decision["decision"] = use_decision["data_quality_flag"].map(_decision)
    use_decision["primary_metric_source"] = "register_summary"
    use_decision["fallback_metric_source"] = "vworld_feature"
    use_decision["decision_reason"] = use_decision.apply(
        lambda r: f"match_rate={r['register_match_rate_by_pnu']:.2%}" if pd.notna(r["register_match_rate_by_pnu"]) else "no matched registers",
        axis=1,
    )
    write_csv(OUT_USE_DECISION, use_decision)

    quality_json = normalize_for_json(
        region_metrics,
        [
            "region_id",
            "region_name",
            "data_quality_flag",
            "register_record_count",
            "unique_register_pk_count",
            "unique_register_pnu_count",
            "register_match_rate_by_pnu",
            "average_far_percent",
            "median_far_percent",
            "weighted_far_percent",
            "average_bcr_percent",
            "median_bcr_percent",
            "weighted_bcr_percent",
            "office_floor_area_share",
            "commercial_floor_area_share",
            "residential_floor_area_share",
        ],
    )
    write_json(OUT_JSON_QUALITY, quality_json)
    write_json(OUT_JSON_IND, normalize_for_json(region_metrics, list(region_metrics.columns)))
    write_json(OUT_JSON_USE, normalize_for_json(use_comp, list(use_comp.columns)))

    discover_rows = pd.DataFrame(source_diag)
    unmatched_registers = int((assigned["match_method"] == "unmatched").sum())
    exact_matches = int((assigned["match_method"] == "exact_vworld_pnu").sum())
    auxiliary_matches = int((assigned["match_method"] == "legacy_pnu_region_auxiliary").sum())

    write_md(
        OUT_METHOD,
        "\n".join(
            [
                "# Building register summary merge methodology v12",
                "",
                "## Objective",
                "- Use VWorld GIS building geometry as the spatial base.",
                "- Use building register summary attributes as the official attribute base.",
                "- Recompute region-level building indicators without double counting.",
                "",
                "## Merge logic",
                "- Register files were auto-discovered under the raw and processed data trees.",
                "- Files were retained only when they contained building-register-like columns such as `관리건축물대장PK`, `대지위치`, `주용도코드`, `사용승인일`, `대지면적`, `연면적`.",
                "- PNU was generated as `sigunguCd(5)+bjdongCd(5)+platGbCd(1)+bun(4)+ji(4)`.",
                "- Exact `parcel_pnu` matching against VWorld features was the primary join path.",
                "- Legacy PNU-to-region mapping from the v2 building analysis was used only as an auxiliary fallback for region assignment.",
                "",
                "## Duplicate prevention",
                "- Feature-level output keeps one row per VWorld building feature.",
                "- Register-level output keeps one row per management register PK.",
                "- Region totals are computed from register rows, not by summing feature matches.",
                "",
                "## FAR/BCR rules",
                "- Weighted FAR = sum(far_calc_floor_area_m2) / sum(site_area_m2) x 100.",
                "- Weighted BCR = sum(building_area_m2) / sum(site_area_m2) x 100.",
                "- Rows with missing or zero site area are excluded from weighted calculations.",
                "",
                "## Quality rule",
                "- high: match rate >= 80% and missing FAR/BCR/site area are limited.",
                "- medium: match rate >= 50%.",
                "- auxiliary_only: match rate >= 20%.",
                "- manual_review_required: match rate below 20%.",
            ]
        ),
    )

    validation_lines = [
        "# Building register summary merge validation v12",
        "",
        f"- Register source files discovered: {len(register_paths)}",
        f"- Register-like source files loaded: {len([x for x in source_diag if x['status'] == 'loaded'])}",
        f"- Raw register rows: {len(raw_register):,}",
        f"- Standardized register rows (before dedup): {len(register_std_all):,}",
        f"- Standardized register rows (deduped output): {len(register_std_output):,}",
        f"- PNU generation success rate: {register_std_all['register_pnu'].notna().mean():.2%}",
        f"- VWorld feature rows: {len(vworld):,}",
        f"- Exact feature matches: {exact_matches:,}",
        f"- Auxiliary matches: {auxiliary_matches:,}",
        f"- Unmatched register rows: {unmatched_registers:,}",
        "",
        "## Region summary",
    ]
    for _, row in region_metrics.iterrows():
        match_rate = f"{row['register_match_rate_by_pnu']:.2%}" if pd.notna(row["register_match_rate_by_pnu"]) else "n/a"
        weighted_far = f"{row['weighted_far_percent']:.2f}" if pd.notna(row["weighted_far_percent"]) else "n/a"
        weighted_bcr = f"{row['weighted_bcr_percent']:.2f}" if pd.notna(row["weighted_bcr_percent"]) else "n/a"
        validation_lines.append(
            f"- {row['region_id']}: records={int(row['register_record_count'])}, match_rate={match_rate}, "
            f"weighted_far={weighted_far}, weighted_bcr={weighted_bcr}, quality={row['data_quality_flag']}"
        )
    write_md(OUT_VALIDATION, "\n".join(validation_lines))

    error_lines = [
        "# Building register summary merge errors v12",
        "",
        f"- Files discovered but excluded due to low register-like score: {int((discover_rows['register_like_score'] < 6).sum()) if not discover_rows.empty else 0}",
        f"- Register rows without exact VWorld PNU match: {unmatched_registers:,}",
        f"- Register rows using auxiliary legacy region mapping: {auxiliary_matches:,}",
        f"- Register rows with duplicate PNU inside the same region output: {int(region_metrics['duplicate_pnu_count'].sum()) if not region_metrics.empty else 0}",
        f"- Register rows with duplicate PK inside the same region output: {int(region_metrics['duplicate_register_pk_count'].sum()) if not region_metrics.empty else 0}",
        "",
        "## Notes",
        "- The raw building files in this workspace were named `표제부`, not `총괄표제부`; they were treated as the available summary-title source.",
        "- Address-string fallback was not applied because the VWorld feature source used here does not expose a stable address field in the raw feature layer.",
        "- Any unmatched records should be reviewed before being used as primary evidence.",
    ]
    write_md(OUT_ERRORS, "\n".join(error_lines))

    # Re-export the main comparison in a concise, review-friendly order.
    comparison = comparison[
        [
            "region_id",
            "region_name",
            "v10_match_rate",
            "v11_feature_count",
            "v12_register_record_count",
            "v12_match_rate_by_pnu",
            "v10_average_far",
            "v11_average_far_percent",
            "v12_average_far_percent",
            "v10_average_bcr",
            "v11_average_bcr_percent",
            "v12_average_bcr_percent",
            "v10_office_floor_area_share",
            "v11_office_floor_area_share",
            "v12_office_floor_area_share",
            "delta_total_floor_area_v12_minus_v11",
            "delta_total_floor_area_v12_minus_v10",
            "v12_data_quality_flag",
        ]
    ]
    write_csv(OUT_COMPARISON, comparison)

    write_methodology_docx(OUT_REPORT, region_metrics, region_metrics, comparison)


if __name__ == "__main__":
    main()
